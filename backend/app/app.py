from fastapi import FastAPI, Request, Form, UploadFile, File, Query, Depends
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI
import os
from dotenv import load_dotenv
from uuid import uuid4
from fastapi.responses import JSONResponse, FileResponse
from docx import Document
from typing import Optional, List
from sqlalchemy.orm import Session
from .init_db import init_db
from .models import ChatSession, Message, Feedback, User, SupportingDocument, SupportingDocumentChunk, SupportingDocumentChunkVector
from .auth_utils import hash_password, parse_user_id, verify_password
from .chat_utils import derive_session_preview, get_chat_history, summarize_old_messages
from .db import get_db
from .file_utils import (
    build_inline_disposition,
    build_view_file_link,
    collect_safe_candidate_paths,
    ensure_filename_with_extension,
    extract_edi_section,
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_xlsx,
    first_existing_path,
    get_last_paragraph_before_edi,
    guess_media_type,
    insert_paragraph_after,
    load_edits,
    paragraph_match_score,
    remove_edi_markers,
    save_edits,
    sorted_prefixed_file_paths,
)
from .prompt_utils import SYSTEM_PROMPT
from .rag_utils import index_supporting_document_chunks, retrieve_supporting_doc_context
from .schemas import AuthRequest, AuthResponse, ForgotPasswordRequest
from fastapi.responses import StreamingResponse
import io
import json
import logging
from pdf2docx import Converter
from starlette.concurrency import run_in_threadpool

#Load OpenAI API key
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
logger = logging.getLogger(__name__)

#Initialize DB
init_db()

#Initialize app
app = FastAPI()

# CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # Vite default
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

#Initialize maximum number of messages in the chat history
MAX_HISTORY = 20

#Create folders to save lesson plans and supporting documents
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LESSON_PLANS_DIR = os.path.join(BASE_DIR, "lessonPlans")
SUPPORTING_DOCS_DIR = os.path.join(BASE_DIR, "supportingDocuments")
os.makedirs(LESSON_PLANS_DIR, exist_ok=True)
os.makedirs(SUPPORTING_DOCS_DIR, exist_ok=True)

#Check write permissions (helps detect permission issues in deployed version)
if not os.access(LESSON_PLANS_DIR, os.W_OK):
    print(f"Warning: Lesson Plans directory may not be writable: {LESSON_PLANS_DIR}")
if not os.access(SUPPORTING_DOCS_DIR, os.W_OK):
    print(f"Warning: Supporting Documents directory may not be writable: {SUPPORTING_DOCS_DIR}")

@app.post("/register", response_model=AuthResponse)
def register(payload: AuthRequest, db: Session = Depends(get_db)):
    username = payload.username.strip()
    password = payload.password
    name = (payload.name or "").strip()
    email = (payload.email or "").strip()
    institution = (payload.institution or "").strip()
    if not username.strip() or not password:
        return JSONResponse(status_code=400, content={"success": False, "message": "Username and password are required."})
    if not name or not email or not institution:
        return JSONResponse(
            status_code=400,
            content={"success": False, "message": "Name, email, and institution are required for registration."},
        )

    existing_user = db.query(User).filter(User.username == username).first()
    if existing_user:
        return JSONResponse(status_code=400, content={"success": False, "message": "Username already exists."})

    password_hash = hash_password(password)
    user = User(username=username, password_hash=password_hash, name=name, email=email, institution=institution)
    db.add(user)
    db.commit()
    db.refresh(user)

    return {"success": True, "user_id": user.id, "username": user.username}

@app.post("/login", response_model=AuthResponse)
def login(payload: AuthRequest, db: Session = Depends(get_db)):
    username = payload.username.strip()    
    password = payload.password
    if not username.strip() or not password:
        return JSONResponse(status_code=400, content={"success": False, "message": "Username and password are required."})

    user = db.query(User).filter(User.username == username).first()
    if not user or not verify_password(password, user.password_hash):
        return JSONResponse(status_code=401, content={"success": False, "message": "Invalid username or password."})

    return {"success": True, "user_id": user.id, "username": user.username}


@app.post("/forgot-password", response_model=AuthResponse)
def forgot_password(payload: ForgotPasswordRequest, db: Session = Depends(get_db)):
    username = payload.username.strip()
    email = payload.email.strip().lower()
    new_password = payload.new_password

    if not username or not email or not new_password:
        return JSONResponse(
            status_code=400,
            content={"success": False, "message": "Username, email, and new password are required."},
        )

    user = db.query(User).filter(User.username == username).first()
    if not user or not user.email or user.email.strip().lower() != email:
        return JSONResponse(
            status_code=400,
            content={"success": False, "message": "The username and email do not match our records."},
        )

    user.password_hash = hash_password(new_password)
    db.commit()

    return {"success": True, "message": "Password reset successful. You can sign in with your new password."}

@app.post("/chatStart")
async def chatStart(user_id: Optional[str] = Form(None), db: Session = Depends(get_db)):
    session_id = str(uuid4())  # Create unique session ID
    user_id_int = parse_user_id(user_id)
    print(f"Starting new chat session: {session_id} for user_id: {user_id_int} (received: {user_id})")
    # Initiate conversation executing system prompt
    db.add(ChatSession(id=session_id, user_id=user_id_int))
    db.add(Message(session_id=session_id, role="system", content=SYSTEM_PROMPT, visible=False))

    try:
        response = await run_in_threadpool(
            client.chat.completions.create,
            model="gpt-4.1-mini",
            messages=[{"role": "system", "content": SYSTEM_PROMPT}],
        )
    except Exception:
        logger.exception("OpenAI call failed while starting chat session %s", session_id)
        return JSONResponse(status_code=502, content={"error": "Failed to start chat. Please try again."})
        
    api_response = response.choices[0].message.content

    db.add(Message(session_id=session_id, role="assistant", content=api_response))
    db.commit()

    return {"response": api_response, "session_id" :session_id}

@app.post("/chatContinue")
async def chatContinue(
    request: Request,
    message: str = Form(None),
    session_id: str = Form(...),
    files: Optional[List[UploadFile]] = File(None),
    file_tuples: str = Form("[]"),
    db: Session = Depends(get_db),
):

    uploaded_files = files or []
    file_type_by_index = {}

    # Parse tuple metadata: [[index, file_type, file_name], ...]
    try:
        parsed_tuples = json.loads(file_tuples or "[]")
        for item in parsed_tuples:
            if not isinstance(item, list) or len(item) < 2:
                continue
            idx = int(item[0])
            current_type = item[1]
            if current_type in ("lesson_plan", "supporting_document"):
                file_type_by_index[idx] = current_type
    except Exception:
        file_type_by_index = {}

    #Retrieve chat session
    chat_session = db.query(ChatSession).filter_by(id=session_id).first()
    if not chat_session:
        return JSONResponse(status_code=404, content={"error": "Session not found."})

    #Generate chat history summary if it is not available 
    if not chat_session.summary:
        try:
            summary = summarize_old_messages(session_id, db, client, MAX_HISTORY)
            if summary:
                chat_session.summary = summary
                db.commit()
        except Exception:
            logger.exception("Failed to summarize old messages for session %s", session_id)

    #Inject chat history and user message to the prompt
    try:
        chat_messages = get_chat_history(session_id, db, SYSTEM_PROMPT, MAX_HISTORY)
    except Exception:
        logger.exception("Failed to build chat history for session %s", session_id)
        return JSONResponse(status_code=500, content={"error": "Failed to build chat context."})

    # Keep separate context buckets so lesson-plan content and supporting-document
    # content can be handled differently downstream.
    uploaded_contexts = []
    uploaded_supporting_docs = []

    for idx, file in enumerate(uploaded_files):
        current_type = file_type_by_index.get(idx, "lesson_plan")
        original_file_name = os.path.splitext(file.filename)[0].strip().rstrip(".") or "uploaded_file"
        original_ext = os.path.splitext(file.filename)[1]
        unique_id = uuid4()
        original_stored_name = f"{unique_id}_{original_file_name}{original_ext}"

        upload_dir = LESSON_PLANS_DIR if current_type == "lesson_plan" else SUPPORTING_DOCS_DIR
        file_path = os.path.abspath(os.path.join(upload_dir, original_stored_name))
        
        try:
            with open(file_path, "wb") as f:
                f.write(file.file.read())
        except Exception:
            logger.exception("Failed to save uploaded file for session %s", session_id)
            return JSONResponse(status_code=500, content={"error": "Failed to save uploaded file."})

        working_file_name = original_stored_name
        working_file_path = file_path
        working_file_type = file.content_type

        # Convert PDF to DOCX for lesson plans only
        if current_type == "lesson_plan" and file.filename.lower().endswith(".pdf"):
            working_file_name = f"{unique_id}_{original_file_name}.docx"
            output_path = os.path.join(upload_dir, working_file_name)
            try:
                cv = Converter(file_path)
                cv.convert(output_path)
                cv.close()
            except Exception:
                logger.exception("Failed PDF to DOCX conversion for session %s", session_id)
                return JSONResponse(status_code=400, content={"error": "Could not process uploaded PDF lesson plan."})
            working_file_path = output_path
            working_file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Extract context from the working file so later edit anchors match the
        # document users will actually download, especially after PDF->DOCX conversion.
        source_path = working_file_path
        if source_path.lower().endswith(".docx"):
            try:
                with open(source_path, "rb") as f:
                    file_content = extract_text_from_docx(f.read())
            except Exception:
                logger.exception("Failed to extract DOCX content for session %s", session_id)
                file_content = "[File uploaded, but content extraction failed.]"
        elif source_path.lower().endswith(".pdf"):
            try:
                with open(source_path, "rb") as f:
                    file_content = extract_text_from_pdf(f.read())
            except Exception:
                logger.exception("Failed to extract PDF content for session %s", session_id)
                file_content = "[File uploaded, but content extraction failed.]"
        elif source_path.lower().endswith(".xlsx"):
            try:
                with open(source_path, "rb") as f:
                    file_content = extract_text_from_xlsx(f.read())
            except Exception:
                logger.exception("Failed to extract XLSX content for session %s", session_id)
                file_content = "[File uploaded, but content extraction failed.]"
        else:
            try:
                with open(source_path, "r", encoding="utf-8") as f:
                    file_content = f.read()
            except Exception:
                file_content = "[File uploaded, but not readable.]"

        if current_type == "lesson_plan":
            uploaded_contexts.append(f"Lesson Plan:\n{file_content}")
            file_link = build_view_file_link(
                request=request,
                session_id=session_id,
                file_kind="lesson_plan",
                file_id=str(unique_id),
            )
            db.add(Message(session_id=session_id, role="user", content=f"📎 [View lesson plan: {file.filename}]", file_link=file_link))
            db.add(Message(session_id=session_id, role="user", content=f"Lesson Plan:\n{file_content}", visible=False))

            # Persist the latest lesson-plan upload as the editable source of truth.
            chat_session.original_lesson = file_content
            chat_session.updated_lesson = file_content
            chat_session.file_name = file.filename
            chat_session.file_path = file_path
            chat_session.file_type = file.content_type
            chat_session.working_file_name = working_file_name
            chat_session.working_file_path = working_file_path
            chat_session.working_file_type = working_file_type
            chat_session.suggested_edits = json.dumps([])
        else:
            uploaded_supporting_docs.append(file.filename)
            file_link = build_view_file_link(
                request=request,
                session_id=session_id,
                file_kind="supporting_document",
                file_id=str(unique_id),
            )
            db.add(Message(session_id=session_id, role="user", content=f"📎 [View supporting document: {file.filename}]", file_link=file_link))

            supporting_doc = SupportingDocument(
                session_id=session_id,
                document_name=file.filename,
                file_path=file_path,
                file_type=file.content_type
            )
            db.add(supporting_doc)
            db.flush()
            chunk_count = index_supporting_document_chunks(db, client, session_id, supporting_doc.id, file_content)
            db.add(
                Message(
                    session_id=session_id,
                    role="system",
                    content=f"Supporting document indexed with embeddings: {file.filename} ({chunk_count} chunks).",
                    visible=False,
                )
            )

    if uploaded_contexts and message:
        chat_messages.append({"role": "user", "content": "\n\n".join(uploaded_contexts) + "\n" + message})
        db.add(Message(session_id=session_id, role="user", content=message))
    elif uploaded_contexts:
        chat_messages.append({"role": "user", "content": "\n\n".join(uploaded_contexts)})
    elif message:
        chat_messages.append({"role": "user", "content": message})
        db.add(Message(session_id=session_id, role="user", content=message))

    retrieval_query = (message or "").strip()
    if retrieval_query and uploaded_supporting_docs:
        retrieval_query = (
            f"{retrieval_query}\n\n"
            "Relevant uploaded supporting document(s): "
            f"{', '.join(uploaded_supporting_docs)}"
        )
    if not retrieval_query and uploaded_supporting_docs:
        retrieval_query = "summarise uploaded supporting documents and ask how to use them"

    rag_context = retrieve_supporting_doc_context(db, client, session_id, retrieval_query)
    if rag_context:
        chat_messages.append(
            {
                "role": "system",
                "content": (
                    "Use the following retrieved supporting-document context when it is relevant to the educator's request. "
                    "\n\n"
                    f"{rag_context}"
                ),
            }
        )
        print(rag_context)

    # If the user only uploaded supporting docs, prompt the model to acknowledge
    # them explicitly instead of treating the turn as silent document ingestion.
    if uploaded_supporting_docs and not message:
        docs_list = ", ".join(uploaded_supporting_docs)
        chat_messages.append(
            {
                "role": "user",
                "content": (
                    "I uploaded supporting document(s): "
                    f"{docs_list}. Please acknowledge and briefly summarize their focus, "
                    "then ask how I want to use them with my lesson plan."
                ),
            }
        )

    try:
        response = await run_in_threadpool(
            client.chat.completions.create,
            model="gpt-4.1-mini",
            messages=chat_messages,
        )
    except Exception:
        logger.exception("OpenAI call failed while continuing session %s", session_id)
        return JSONResponse(status_code=502, content={"error": "Failed to generate assistant response."})
    api_response = response.choices[0].message.content

    db.add(Message(session_id=session_id, role="assistant", content=api_response))
    db.commit()

    return {"response": api_response, "session_id": session_id}


#Retrieve chat sessions for chat history   
@app.get("/sessions")
def get_sessions(user_id: Optional[str] = Query(None), db: Session = Depends(get_db)):
        query = db.query(ChatSession)
        user_id_int = parse_user_id(user_id)
        print(f"Fetching sessions for user_id: {user_id_int} (received: {user_id})")
        if user_id_int is not None:
            query = query.filter(ChatSession.user_id == user_id_int)

        sessions = query.order_by(ChatSession.created_at.desc()).all()
        results = []
        for s in sessions:
            preview = derive_session_preview(s.id, db)
            results.append({
                "id": s.id,
                "created_at": s.created_at.isoformat() if s.created_at else None,
                "lesson_preview": preview,
            })
        return JSONResponse(content=results)

#Retrieve messages of the selected chat session from the chat history
@app.get("/sessionMessages")
def get_session_messages(session_id: str = Query(...), db: Session = Depends(get_db)):
    file = ""
    query = db.query(Message).filter_by(session_id=session_id)
    query = query.filter_by(visible=True)
    messages = query.order_by(Message.timestamp).all()
    session = db.query(ChatSession).filter_by(id=session_id).first()
    if session:
        if session.updated_lesson:
            file = session.updated_lesson
        else:
            file = session.original_lesson
            
    results = [{"role": m.role, "content": m.content, "file_link": m.file_link} for m in messages]
    return {"file": file, "messages": results}

#Retrieve supporting documents for a session
@app.get("/supportingDocs")
def get_supporting_docs(session_id: str = Query(...), db: Session = Depends(get_db)):
    supporting_docs = db.query(SupportingDocument).filter_by(session_id=session_id).order_by(SupportingDocument.uploaded_at.desc()).all()
    results = [
        {
            "id": doc.id,
            "document_name": doc.document_name,
            "file_path": doc.file_path,
            "uploaded_at": doc.uploaded_at.isoformat() if doc.uploaded_at else None,
        }
        for doc in supporting_docs
    ]
    return {"supporting_docs": results}

#Remove a supporting document
@app.post("/removeSupportingDoc")
def remove_supporting_doc(doc_id: int = Form(...), db: Session = Depends(get_db)):
    supporting_doc = db.query(SupportingDocument).filter_by(id=doc_id).first()
    
    if not supporting_doc:
        return JSONResponse(status_code=404, content={"error": "Document not found."})
    
    # Delete the file from storage
    try:
        if os.path.exists(supporting_doc.file_path):
            os.remove(supporting_doc.file_path)
    except Exception as e:
        print(f"Warning: Could not delete file {supporting_doc.file_path}: {str(e)}")
    
    # Delete retrieval vectors/chunks and then metadata from database
    db.query(SupportingDocumentChunkVector).filter_by(supporting_document_id=supporting_doc.id).delete()
    db.query(SupportingDocumentChunk).filter_by(supporting_document_id=supporting_doc.id).delete()
    db.delete(supporting_doc)
    db.commit()
    
    return {"success": True, "message": "Supporting document removed."}



#View uploaded lesson plan in chat history
@app.get("/viewFile", name="view_file")
def view_file(
    session_id: str = Query(...),
    file_id: Optional[str] = Query(None),
    stored_name: Optional[str] = Query(None),
    file_kind: str = Query("lesson_plan"),
    display_name: Optional[str] = Query(None),
    db: Session = Depends(get_db),
):
    chat_session = db.query(ChatSession).filter_by(id=session_id).first()

    # Preferred behavior for new links: resolve by immutable file UUID so a link
    # never drifts to a later upload with the same session.
    if file_id:
        safe_id = (file_id or "").strip()
        base_dir = LESSON_PLANS_DIR if file_kind == "lesson_plan" else SUPPORTING_DOCS_DIR
        base_dir_abs = os.path.abspath(base_dir)

        preferred_exts = [".pdf", ".docx", ".txt", ".csv", ".xlsx", ".xls", ".json"]
        candidate_paths = sorted_prefixed_file_paths(
            base_dir_abs,
            f"{safe_id}_",
            preferred_exts,
        )
        resolved_path = first_existing_path(candidate_paths)
        if not resolved_path:
            return JSONResponse(status_code=404, content={"error": "File missing on server."})

        media_type = guess_media_type(resolved_path)

        derived_name = os.path.basename(resolved_path).split("_", 1)[1] if "_" in os.path.basename(resolved_path) else os.path.basename(resolved_path)
        response_name = ensure_filename_with_extension(display_name or derived_name, resolved_path)
        resolved_ext = os.path.splitext(resolved_path)[1]
        name_root, name_ext = os.path.splitext(response_name)
        if resolved_ext and name_ext.lower() != resolved_ext.lower():
            response_name = f"{name_root}{resolved_ext}"

        return FileResponse(
            path=resolved_path,
            media_type=media_type,
            filename=response_name,
            headers={
                "Content-Disposition": build_inline_disposition(response_name)
            }
        )

    # Backward-compatible behavior: resolve old links by stored_name.
    if stored_name:
        safe_name = os.path.basename(stored_name)
        sanitized_name = safe_name.strip().rstrip(".")
        if safe_name != stored_name:
            return JSONResponse(status_code=400, content={"error": "Invalid file reference."})

        base_dir = LESSON_PLANS_DIR if file_kind == "lesson_plan" else SUPPORTING_DOCS_DIR
        base_dir_abs = os.path.abspath(base_dir)
        candidate_names = [safe_name]
        if sanitized_name and sanitized_name not in candidate_names:
            candidate_names.append(sanitized_name)

        candidate_paths = collect_safe_candidate_paths(base_dir_abs, candidate_names)

        # If exact name lookup fails (e.g. trailing-space normalization differences),
        # resolve by immutable UUID prefix in the same folder only.
        uuid_prefix = sanitized_name.split("_", 1)[0] if "_" in sanitized_name else ""
        expected_ext = os.path.splitext(sanitized_name)[1].lower()
        display_ext = os.path.splitext((display_name or "").strip())[1].lower()
        if uuid_prefix:
            preferred_exts = []
            if expected_ext:
                preferred_exts.append(expected_ext)
            if display_ext and display_ext not in preferred_exts:
                preferred_exts.append(display_ext)
            if file_kind == "lesson_plan":
                for ext in [".pdf", ".docx"]:
                    if ext not in preferred_exts:
                        preferred_exts.append(ext)

            candidate_paths.extend(
                sorted_prefixed_file_paths(
                    base_dir_abs,
                    f"{uuid_prefix}_",
                    preferred_exts,
                )
            )

        resolved_path = first_existing_path(candidate_paths)
        if not resolved_path:
            return JSONResponse(status_code=404, content={"error": "File missing on server."})

        media_type = guess_media_type(resolved_path)

        response_name = ensure_filename_with_extension(display_name or safe_name, resolved_path)
        resolved_ext = os.path.splitext(resolved_path)[1]
        name_root, name_ext = os.path.splitext(response_name)
        if resolved_ext and name_ext.lower() != resolved_ext.lower():
            response_name = f"{name_root}{resolved_ext}"

        return FileResponse(
            path=resolved_path,
            media_type=media_type,
            filename=response_name,
            headers={
                "Content-Disposition": build_inline_disposition(response_name)
            }
        )

    # Legacy behavior: session-level file for old links without stored_name.
    if not chat_session or not chat_session.file_path:
        return JSONResponse(status_code=404, content={"error": "File not found."})
    
    # Safely verify file exists before sending
    if not os.path.exists(chat_session.file_path):
        return JSONResponse(status_code=404, content={"error": "File missing on server."})

    legacy_name = ensure_filename_with_extension(chat_session.file_name, chat_session.file_path)
    legacy_ext = os.path.splitext(chat_session.file_path)[1]
    legacy_root, legacy_name_ext = os.path.splitext(legacy_name)
    if legacy_ext and legacy_name_ext.lower() != legacy_ext.lower():
        legacy_name = f"{legacy_root}{legacy_ext}"

    return FileResponse(
        path=chat_session.file_path,
        media_type=guess_media_type(chat_session.file_path),
        filename=legacy_name,
        headers={
            "Content-Disposition": build_inline_disposition(legacy_name)
        }
    )


#Lesson plan update functionality
@app.post("/updateLesson")
async def update_lesson(session_id: str = Form(...), new_content: str = Form(...), db: Session = Depends(get_db)):
    chat_session = db.query(ChatSession).filter_by(id=session_id).first()
    existing_edits = []
    
    if not chat_session:
        return {"error": "Session not found"}
    
    if chat_session.updated_lesson:
        currentContent = chat_session.updated_lesson
    else:
        currentContent = chat_session.original_lesson
    
    if chat_session.suggested_edits:
        existing_edits = load_edits(chat_session)
    
    if not chat_session.summary:
        try:
            summary = summarize_old_messages(session_id, db, client, MAX_HISTORY)
            if summary:
                chat_session.summary = summary
                db.commit()
        except Exception:
            logger.exception("Failed to summarize old messages for lesson update session %s", session_id)

    #Update lesson plan by appending suggested content using LLM API
    try:
        chat_messages = get_chat_history(session_id, db, SYSTEM_PROMPT, MAX_HISTORY)
    except Exception:
        logger.exception("Failed to build chat history for lesson update session %s", session_id)
        return JSONResponse(status_code=500, content={"error": "Failed to build lesson update context."})
    chat_messages.append({"role": "user", "content": f'''Update the lesson plan by integrating the new content - \n{new_content} in to the current lesson plan - \n{currentContent} appropriately preserving the pedagogical flow. 
                          In the response provide the full content of the updated lesson plan. Do not include any additional texts in the response.
                          When adding new content start with "### EDI integration start.". Mention this in a new line.
                          At the end of the new content mention "### EDI integration end." '''
                          })
    try:
        response = await run_in_threadpool(
            client.chat.completions.create,
            model="gpt-4.1-mini",
            messages=chat_messages,
        )
    except Exception:
        logger.exception("OpenAI call failed for lesson update session %s", session_id)
        return JSONResponse(status_code=502, content={"error": "Failed to generate updated lesson plan."})
    api_response = response.choices[0].message.content
    new_edit = extract_edi_section(api_response)
    target_text = get_last_paragraph_before_edi(api_response)
    clean_api_response = remove_edi_markers(api_response)

    edit = {
        "target_text": target_text,
        "new_content": new_edit
    }
    existing_edits.append(edit)
    save_edits(chat_session, existing_edits)
    chat_session.updated_lesson = clean_api_response #Update updated lesson plan in db 
    success_message = "Lesson plan updated successfully. The updated lesson preview is now available."
    
    full_update_message = f"{success_message}\n\n{clean_api_response}"
    db.add(Message(session_id=session_id, role="assistant", content=full_update_message))
    download_message = "You can download the updated lesson plan by clicking the “Download Updated Lesson Plan” button in the right panel. \n\n Would you like further support with this lesson plan? You can request additional support by selecting a support option from the right panel or by describing your requirements directly in the chat. \n\n If you would like to integrate EDI principles into a different lesson plan, you can upload a new lesson plan at any time."    
    db.add(Message(session_id=session_id, role="assistant", content=download_message))
    
    db.commit()

    return {"response": full_update_message, "download_message": download_message, "session_id": session_id}

@app.get("/previewLesson")
def preview_lesson(session_id: str, db: Session = Depends(get_db)):
    chat_session = db.query(ChatSession).filter_by(id=session_id).first()
    if not chat_session:
        return JSONResponse(status_code=404, content={"error": "Session not found."})

    if not chat_session.updated_lesson:
        return JSONResponse(status_code=404, content={"error": "Updated lesson not found."})
  
    updated = chat_session.updated_lesson

    # HTML preview
    html = "<div style='font-family:Arial;line-height:1.6'>"
    for line in updated.split("\n"):
        if "💡EDI Content:" in line:
            html += f"<div style='background:#fff3cd;padding:8px'>{line}</div>"
        else:
            html += f"<p>{line}</p>"
    html += "</div>"

    bot_message = "You can download the updated lesson plan by clicking the “Download Lesson Plan” button in the right panel. \n\n Would you like further support with this lesson plan? You can request additional support by selecting a support option from the right panel or by describing your requirements directly in the chat. \n\n If you would like to integrate EDI principles into a different lesson plan, you can upload a new lesson plan at any time."
    return {"html": html, "bot_message":bot_message, "session_id": session_id}

#Download updated lesson plan functionality
@app.get("/downloadLesson")
def download_lesson(session_id: str = Query(...), db: Session = Depends(get_db)):
    chat_session = db.query(ChatSession).filter_by(id=session_id).first()
    
    if not chat_session or not chat_session.working_file_path:
        return JSONResponse(status_code=404, content={"error": "Working file path not found."})

    # Create a .docx document
    try:
        doc = Document(chat_session.working_file_path)
    except Exception:
        logger.exception("Failed to load working lesson document for session %s", session_id)
        return JSONResponse(status_code=500, content={"error": "Failed to prepare lesson download."})
    edits = load_edits(chat_session)

    # Keep only the latest edit for each anchor so repeated updates do not stack
    deduped_edits = []
    seen_targets = set()
    for edit in reversed(edits):
        target_text = edit.get("target_text")
        if not target_text or target_text in seen_targets:
            continue
        seen_targets.add(target_text)
        deduped_edits.append(edit)
    deduped_edits.reverse()

    if deduped_edits:
        for e in deduped_edits:
            target_text = e.get("target_text", "")
            if not target_text:
                continue

            best_paragraph = None
            best_score = 0.0
            best_index = -1
            for para_index, para in enumerate(doc.paragraphs):
                score = paragraph_match_score(target_text, para.text)

                if score > best_score or (score == best_score and para_index > best_index):
                    best_score = score
                    best_paragraph = para
                    best_index = para_index

            if best_paragraph and best_score >= 0.3:
                insert_paragraph_after(
                    best_paragraph,
                    f"💡 EDI Content: {e['new_content']}"
                )

    # Save to in-memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=updated_lesson_{session_id[:8]}.docx"
        }
    )

@app.post("/submitFeedback")
def submit_feedback(
    session_id: str = Form(...),
    feedback: str = Form(...),
    user_id: str = Form(...),
    db: Session = Depends(get_db),
):
    db.add(Feedback(session_id=session_id, feedback=feedback, user_id=user_id))
    db.commit()
    return {"message": "Feedback submitted successfully."}