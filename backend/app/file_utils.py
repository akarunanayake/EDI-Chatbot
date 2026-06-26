from io import BytesIO
from difflib import SequenceMatcher
import json
import re
import os
from urllib.parse import quote

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
import fitz
from openpyxl import load_workbook


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    return "\n".join(para.text for para in doc.paragraphs)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page in pdf:
            text += page.get_text()
    return text


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    workbook = load_workbook(filename=BytesIO(file_bytes), data_only=True, read_only=True)
    lines = []

    for sheet in workbook.worksheets:
        lines.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                lines.append("\t".join(cells))

    workbook.close()
    return "\n".join(lines).strip()


def load_edits(session):
    try:
        return json.loads(session.suggested_edits or "[]")
    except Exception:
        return []


def save_edits(session, edits):
    session.suggested_edits = json.dumps(edits)


def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)

    new_para = Paragraph(new_p, paragraph._parent)

    if text:
        run = new_para.add_run(text)
        run.bold = True
        run.italic = False
        run.font.color.rgb = RGBColor(0, 51, 102)

    return new_para


def extract_edi_section(text: str):
    match = re.search(
        r"### EDI integration start\.\s*(.*?)\s*### EDI integration end\.",
        text,
        re.DOTALL,
    )

    if match:
        return match.group(1).strip()
    return None


def get_last_sentence_before_edi(text: str):
    marker = "### EDI integration start."

    if marker not in text:
        return None

    before_edi = text.split(marker)[0].strip()
    sentences = re.split(r"(?<=[.!?])\s+", before_edi)

    for sentence in reversed(sentences):
        if sentence.strip():
            return sentence.strip()

    return None


def get_last_paragraph_before_edi(text: str):
    marker = "### EDI integration start."

    if marker not in text:
        return None

    before = text.split(marker)[0].strip()
    paragraphs = [paragraph.strip() for paragraph in before.split("\n") if paragraph.strip()]

    return paragraphs[-1] if paragraphs else None


def remove_edi_markers(text: str):
    text = text.replace("### EDI integration start.", "")
    text = text.replace("### EDI integration end.", "")
    return text


def normalize_for_match(text: str) -> str:
    # Collapse whitespace and punctuation noise to improve anchor matching.
    return re.sub(r"\W+", " ", (text or "").lower()).strip()


def paragraph_match_score(target_text: str, paragraph_text: str) -> float:
    normalized_target = normalize_for_match(target_text)
    normalized_paragraph = normalize_for_match(paragraph_text)

    if not normalized_target or not normalized_paragraph:
        return 0.0

    target_raw = (target_text or "").strip()
    paragraph_raw = (paragraph_text or "").strip()

    if target_raw and target_raw == paragraph_raw:
        return 1.0

    if target_raw and target_raw in paragraph_raw:
        return 0.98

    if paragraph_raw and paragraph_raw in target_raw:
        return 0.96

    sequence_ratio = SequenceMatcher(None, normalized_target, normalized_paragraph).ratio()
    target_tokens = set(normalized_target.split())
    paragraph_tokens = set(normalized_paragraph.split())

    if not target_tokens or not paragraph_tokens:
        return sequence_ratio

    shared_tokens = target_tokens & paragraph_tokens
    token_overlap = len(shared_tokens) / max(len(target_tokens), 1)

    return max(sequence_ratio, token_overlap)


def ensure_filename_with_extension(name: str | None, file_path: str | None) -> str:
    cleaned_name = (name or "file").replace("\r", "").replace("\n", "").strip()
    if not cleaned_name:
        cleaned_name = "file"

    name_ext = os.path.splitext(cleaned_name)[1]
    path_ext = os.path.splitext(file_path or "")[1]
    if not name_ext and path_ext:
        cleaned_name = f"{cleaned_name}{path_ext}"

    return cleaned_name


def build_inline_disposition(name: str | None) -> str:
    cleaned = (name or "file").replace("\r", "").replace("\n", "").strip()
    if not cleaned:
        cleaned = "file"
    ascii_fallback = cleaned.encode("ascii", "ignore").decode("ascii").strip() or "file"
    ascii_fallback = ascii_fallback.replace('"', "")
    encoded = quote(cleaned)
    return f"inline; filename=\"{ascii_fallback}\"; filename*=UTF-8''{encoded}"


def guess_media_type(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".csv": "text/csv",
        ".xls": "application/vnd.ms-excel",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".json": "application/json",
    }.get(ext, "application/octet-stream")


def collect_safe_candidate_paths(base_dir_abs: str, candidate_names: list[str]) -> list[str]:
    candidate_paths = []
    for candidate_name in candidate_names:
        requested_path = os.path.abspath(os.path.join(base_dir_abs, candidate_name))
        if requested_path.startswith(base_dir_abs + os.sep):
            candidate_paths.append(requested_path)
    return candidate_paths


def sorted_prefixed_file_paths(
    base_dir_abs: str,
    prefix: str,
    preferred_exts: list[str] | None = None,
) -> list[str]:
    try:
        folder_files = os.listdir(base_dir_abs)
    except Exception:
        return []

    matches = [
        name for name in folder_files
        if name.startswith(prefix)
    ]

    ranked_exts = [ext.lower() for ext in (preferred_exts or [])]

    def ext_rank(file_name: str) -> int:
        file_ext = os.path.splitext(file_name)[1].lower()
        if file_ext in ranked_exts:
            return ranked_exts.index(file_ext)
        return len(ranked_exts)

    ranked_matches = sorted(matches, key=lambda name: (ext_rank(name), name))
    return collect_safe_candidate_paths(base_dir_abs, ranked_matches)


def first_existing_path(candidate_paths: list[str]) -> str | None:
    return next((path for path in candidate_paths if os.path.exists(path)), None)
